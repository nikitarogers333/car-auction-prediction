#!/usr/bin/env python3
"""Update writeup_anton_final_v2.docx with H4 ablation results."""

import sys
from copy import deepcopy
from pathlib import Path

import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

INPUT = Path.home() / "Downloads" / "writeup_anton_final_v2.docx"
OUTPUT_DOCS = Path(__file__).resolve().parent.parent / "docs" / "writeup_anton_final_v3.docx"
OUTPUT_DL = Path.home() / "Downloads" / "writeup_anton_final_v3.docx"


def _clear_cell_text(tc_elem):
    """Remove all runs from all paragraphs in a table cell XML element."""
    for p in tc_elem.findall(qn("w:p")):
        for r in p.findall(qn("w:r")):
            p.remove(r)


def _set_cell_text(tc_elem, text, copy_rpr_from=None):
    """Set text in a table cell XML element, optionally copying run formatting."""
    _clear_cell_text(tc_elem)
    p = tc_elem.find(qn("w:p"))
    if p is None:
        p = OxmlElement("w:p")
        tc_elem.append(p)
    for i, line in enumerate(text.split("\n")):
        if i > 0:
            run_br = OxmlElement("w:r")
            run_br.append(OxmlElement("w:br"))
            p.append(run_br)
        run = OxmlElement("w:r")
        if copy_rpr_from is not None:
            rpr = copy_rpr_from.find(qn("w:rPr"))
            if rpr is not None:
                run.append(deepcopy(rpr))
        t = OxmlElement("w:t")
        t.text = line
        t.set(qn("xml:space"), "preserve")
        run.append(t)
        p.append(run)


def _insert_paragraph_after(ref_para, text, bold=False):
    """Insert a new paragraph after ref_para. Returns the new OxmlElement."""
    new_p = OxmlElement("w:p")
    run = OxmlElement("w:r")
    if bold:
        rpr = OxmlElement("w:rPr")
        rpr.append(OxmlElement("w:b"))
        run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    run.append(t)
    new_p.append(run)
    ref_para._element.addnext(new_p)
    return new_p


def main():
    if not INPUT.exists():
        print(f"ERROR: {INPUT} not found"); sys.exit(1)

    doc = docx.Document(str(INPUT))

    # =========================================================
    # 1. Table 1 (Hypotheses): Update H4 row
    # =========================================================
    h_table = doc.tables[1]
    h_table.cell(4, 2).text = "Supported"
    h_table.cell(4, 3).text = "E5: $2,818  A\u2019: $8,045"
    print("  Updated H4 in hypotheses table")

    # =========================================================
    # 2. Table 2 (Results): Insert A' column after E5
    # =========================================================
    results_tbl = doc.tables[2]._tbl
    grid = results_tbl.find(qn("w:tblGrid"))
    grid_cols = grid.findall(qn("w:gridCol"))
    grid_cols[2].addnext(deepcopy(grid_cols[2]))

    aprime_col = [
        "A\u2019: Ablation\nFree-form +\nformula",
        "100%",
        "Never",
        "$8,045",
        "20%",
        "85%",
    ]

    for i, tr in enumerate(results_tbl.findall(qn("w:tr"))):
        tcs = tr.findall(qn("w:tc"))
        new_tc = deepcopy(tcs[2])
        orig_run = tcs[2].find(".//" + qn("w:r"))
        _set_cell_text(new_tc, aprime_col[i] if i < len(aprime_col) else "", copy_rpr_from=orig_run)
        tcs[2].addnext(new_tc)

    print("  Added A' column to results table")

    # =========================================================
    # 3. Table 2: Add "Price varies run-to-run" row
    # =========================================================
    all_tr = results_tbl.findall(qn("w:tr"))
    last_tr = all_tr[-1]
    cv_row = deepcopy(last_tr)
    # Now columns are: Metric, E3, E5, A', RF, XGB (indices 0-5)
    cv_data = [
        "Price varies run-to-run",
        "5%",
        "3.37%",
        "67.71%",
        "0% *",
        "0% *",
    ]
    orig_run = all_tr[1].findall(qn("w:tc"))[0].find(".//" + qn("w:r"))
    for j, tc in enumerate(cv_row.findall(qn("w:tc"))):
        _set_cell_text(tc, cv_data[j] if j < len(cv_data) else "", copy_rpr_from=orig_run)
    last_tr.addprevious(cv_row)
    print("  Added Price CV row")

    # =========================================================
    # 4. Insert H4 explanation after H1 explanation
    # =========================================================
    paras = doc.paragraphs
    # Find "H1 failed" explanation paragraph (the one after the H1 heading)
    target_idx = None
    for i, p in enumerate(paras):
        if "The core finding" in p.text:
            target_idx = i - 1
            break

    if target_idx is not None:
        anchor = paras[target_idx]
        h4_body = (
            "E5 MAE $2,818 vs A\u2019 MAE $8,045 \u2014 nearly 3x worse with the same "
            "formula and same cars. A\u2019 had a 100% valid rate and zero retries \u2014 "
            "the parser never failed. It extracted the wrong values because the LLM\u2019s "
            "free-form prose was ambiguous. Without a schema forcing a decision, the LLM "
            "produced \u201cfairly strong\u201d one run and \u201cmoderate\u201d the next, "
            "which the parser mapped to different values. The 67.71% price CV on A\u2019 "
            "vs 3.37% on E5 is the proof: the schema forces the LLM to commit to a "
            "specific interpretation of the car at the reasoning step, not just the "
            "formatting step."
        )
        h4_title = "H4: Schema enforcement is a precision constraint on reasoning, not just formatting"

        # Insert body first (addnext puts it right after), then heading (pushes body down)
        _insert_paragraph_after(anchor, h4_body)
        _insert_paragraph_after(anchor, "")  # blank line
        _insert_paragraph_after(anchor, h4_title, bold=True)
        _insert_paragraph_after(anchor, "")  # blank line before heading
        print(f"  Inserted H4 explanation after paragraph {target_idx}")
    else:
        print("  WARNING: Could not find insertion point for H4 explanation")

    # =========================================================
    # 5. Update Next Steps section
    # =========================================================
    paras = doc.paragraphs  # re-enumerate after insertions
    for p in paras:
        if "H4 ablation" in p.text and "does structured" in p.text:
            for run in p.runs:
                run.text = ""
            if p.runs:
                p.runs[0].text = (
                    "H4 ablation: Complete \u2014 Supported. Schema enforcement "
                    "improves accuracy 3x over free-form extraction with the same formula."
                )
            print("  Updated H4 next-step bullet")

        if "H4 requires building" in p.text:
            for run in p.runs:
                run.text = ""
            if p.runs:
                p.runs[0].text = (
                    "H4 has been completed. The A\u2019 pipeline (free-form LLM + parser + "
                    "same formula) ran at the same settings. Result: E5 MAE $2,818 vs "
                    "A\u2019 MAE $8,045 \u2014 schema enforcement matters."
                )
            print("  Updated H4 explanation bullet")

    # =========================================================
    # 6. Update Step 2 description to mention 5 pipelines
    # =========================================================
    for p in paras:
        if "All four pipelines run simultaneously" in p.text:
            for run in p.runs:
                run.text = run.text.replace("four pipelines", "five pipelines")
            print("  Updated pipeline count to five")

    # =========================================================
    # Save
    # =========================================================
    OUTPUT_DOCS.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_DOCS))
    doc.save(str(OUTPUT_DL))
    print(f"\n  Saved to {OUTPUT_DOCS}")
    print(f"  Saved to {OUTPUT_DL}")


if __name__ == "__main__":
    main()
